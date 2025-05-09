import matplotlib.pyplot as plt
import numpy as np
import openpyxl
import pandas as pd
import os
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH# 导入 WD_ALIGN_PARAGRAPH 枚举，用于设置段落对齐方式
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Pt
from docx.oxml.ns import qn #用于生成带有命名空间的 XML 标签


"""
功能：将输入的单元格边框设置为不同格式
elm: 要设置边框的 XML 元素。
        border_name: 边框名称（例如 'top', 'left', 'bottom', 'right'）。
        sz: 边框粗细（例如 '4' 或 '8'，表示磅值）。
        color: 边框颜色（例如 '000000' 表示黑色）。
        val: 边框样式（例如 'single' 表示单实线，'double' 表示双实线）。
"""
def set_border(elm,border_name,sz,color,val):
    border = elm.find(qn(f'w:{border_name}'))
    if border is None:
        border = elm.makeelement(qn(f'w:{border_name}'))
        elm.append(border)
    
    border.set(qn('w:sz'),str(sz))
    border.set(qn('w:color'),color)
    border.set(qn('w:val'),val)


sys.stdout.reconfigure(encoding='utf-8') #用于保证OUTPUT输出中文字符不乱码

## 功能：将Table的外边框设置为双实线 ##
def table_border(table):
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    for i in range(num_rows):
        for j in range(num_cols):
            cell = table.cell(i,j)
            tc = cell._element
            # 检查tcBorders是否存在
            tcBorders_list = tc.xpath('.//w:tcBorders')
            if not tcBorders_list:
                # 如果不存在，创建tcBorders元素
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            else:
                tcBorders = tcBorders_list[0]

            #处理顶部边框
            if i == 0:
                set_border(tcBorders,'top',4,'000000','double')
                # top_inner = tcBorders.makeelement(qn('w:top'),{
                #     qn('w:sz'):'4',
                #     qn('w:color'):'000000',
                #     qn('w:val'):'double'
                # }
                # )
                # tcBorders.append(top_inner)
            else:
                set_border(tcBorders,'top',8,'000000','single')
                # #非顶部单元格
                # top_border = tcBorders.find(qn('w:top'))
                # if top_border is not None:
                #     tcBorders.remove(top_border)
        

            #处理左侧边框
            if j == 0:
                set_border(tcBorders,'left',4,'000000','double')
            else:
                set_border(tcBorders,'left',8,'000000','single')
            
            #处理右侧边框
            if j == num_cols-1:
                set_border(tcBorders,'right',4,'000000','double')
            else:
                set_border(tcBorders,'right',8,'000000','single')

            #处理底部边框
            if i == num_rows-1:
                set_border(tcBorders,'bottom',4,'000000','double')
            else:
                set_border(tcBorders,'bottom',8,'000000','single')

### Excel 转 Word ###
def append_excel_to_word_table(excel_file, word_file, sheet_name, start_row, end_row, start_col, end_col):
    """
    将 Excel 中指定行列的数据追加到 Word 表格中。

    参数：
    excel_file (str): Excel 文件路径。
    word_file (str): Word 文件保存路径。
    sheet_name (str): Excel 工作表名称。
    start_row (int): 起始行号（从 0 开始）。
    end_row (int): 结束行号（不包含）。
    start_col (int): 起始列号（从 0 开始）。
    end_col (int): 结束列号（不包含）。
    """
    try:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
    except FileNotFoundError:
        print(f"找不到文件：{excel_file}")
        return
    except Exception as e:
        print(f"读取 Excel 文件时出错：{e}")
        return

    # 提取指定行列的数据
    selected_data = df.iloc[start_row:end_row, start_col:end_col] # 使用 iloc 提取指定范围的数据

    # 检查 Word 文件是否存在，如果存在则读取，否则创建新文档
    if os.path.exists(word_file):
        doc = Document(word_file)
    else:
        doc = Document()

    if sheet_name == 'TIEM':
        image_path = f'#XTHF_Time_{M}M.png'
        doc.add_picture(image_path,height=Cm(10))
    elif sheet_name == 'FRE1':
        image_path = f'#XTHF_FRE_{M}M.png'
        doc.add_picture(image_path,height=Cm(10))

    # 添加空行
    doc.add_paragraph("")  # 添加一个空段落
    # 添加新的表格   先行+2：添加标题和条件 后列
    table = doc.add_table(rows=selected_data.shape[0] + 1 + 1, cols=selected_data.shape[1])

    chinese_font = '宋体'
    english_font = 'Times New Roman'

    table.cell(0,1).merge(table.cell(0,5))#合并第1行的后面列
    table.style = 'Table Grid' #表格样式为网格

    table_border(table)
    cell1 = table.cell(0,1)
    parag = cell1.paragraphs[0]

    for run in parag.runs:
        run.clear()
    # 段落中添加一个运行对象，并写入文本
    if sheet_name == 'TIEM':
        if M!= 5:
            run = parag.add_run(f"启动时间(us) @{CryFre[M]}M晶振")
        else:
            run = parag.add_run(f"启动时间(us) @{CryFre[M]}M陶振")
    elif sheet_name == 'FRE1':
        if M!= 5:
            run = parag.add_run(f"振荡频率(kHz) @{CryFre[M]}M晶振")
        else:
            run = parag.add_run(f"振荡频率(kHz) @{CryFre[M]}M陶振")
    

    # 设置英文字体
    run.font.name = english_font
    # 设置中文字体
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'),chinese_font)

    run.font.size = Pt(12)

    table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #设置表头居中对其
    # table.cell(1,1).merge(table.cell(1,5))#合并第2行的后面列

    # 写入表头（使用 Excel 的列索引）
    for j, col_index in enumerate(range(start_col, end_col)): #遍历行
        cell = table.cell(0 + 1, j) #遍历列
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = str(df.columns[col_index])  # 获取原始列名，写入单元格
        cell.paragraphs[0].runs[0].font.size = Pt(10) #设置表头字体大小
        cell.paragraphs[0].runs[0].font.name = english_font
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #设置表头居中对其
        

    # 写入数据
    for i, row in enumerate(range(start_row, end_row)): #遍历行
        for j, col_index in enumerate(range(start_col, end_col)):#遍历列
            cell = table.cell(i + 1 + 1, j) #获取单元格
            cell.text = str(df.iloc[row, col_index]) #写入数据
            cell.paragraphs[0].runs[0].font.size = Pt(10) #设置数据字体大小
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #设置表头居中对其            cell.paragraphs[0].runs[0].font.name = english_font

    cell2 = table.cell(1, 0)

    cell2.paragraphs[0].clear()

    paragraph1 = cell2.paragraphs[0]
    run1 = paragraph1.add_run("芯片编号")
    run1.font.name = english_font
    # 设置中文字体
    r1 = run1._element
    r1.rPr.rFonts.set(qn('w:eastAsia'),chinese_font)
    paragraph1.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # run1.font.size = Pt(12)

    parag2 = cell2.add_paragraph()
    run2 = parag2.add_run("振荡强度")
    parag2.alignment = WD_TABLE_ALIGNMENT.LEFT
    run2.font.size = Pt(10)
    run2.font.name = "宋体"
    # 设置中文字体
    r2 = run2._element
    r2.rPr.rFonts.set(qn('w:eastAsia'),chinese_font)

    try:
        doc.save(word_file)
        print(f"成功将 Excel 数据追加到 Word 表格：{word_file}")
    except Exception as e:
        print(f"保存 Word 文件时出错：{e}")

### 画图数据设置 ###
# 设置matplotlib的字体
plt.rcParams['font.sans-serif'] = ['Times New Roman']  # 指定默认字体
plt.rcParams['font.sans-serif'].insert(0,'SimSun')
plt.rcParams['axes.unicode_minus'] = False  # 解决保存图像是负号'-'显示为方块的问题

plt.rcParams['mathtext.fontset'] = 'cm'
plt.rcParams['mathtext.rm'] = 'Times New Roman'
plt.rcParams['mathtext.it'] = 'serif:italic'
plt.rcParams['mathtext.bf'] = 'serif:bold'

def substract_column(array,col_num):
    #复制原数组，避免修改原数组
    new_array = array.copy()
    #获取中间的数据
    eight_column = array[:,col_num]
    #将第四列的数值设为0
    new_array[:,col_num] = 0
    #遍历其他列
    for i in range(array.shape[1]):
        if i!=col_num:
            #每一列减去第八列数据
            new_array[:,i] -= eight_column
            new_array[:,i] /= eight_column
    return new_array*1e6

def transform_array(arr, subtract, divide):
    """
    减去数组中的每个元素一个数，然后除以另一个数。

    :param arr: NumPy数组
    :param subtract: 要从每个元素中减去的数
    :param divide: 要除以的数
    :return: 修改后的NumPy数组
    """
    if divide == 0:
        raise ValueError("除数不能为0")
    
    return ((arr - subtract) / divide)*1e6

#Crystal Frequency
M_Fre = [0,1,2,3,4,5]
CryFre = [8,16,24,32,40,8] #前面是晶振 最后一个是陶振
#VDD
VDD = [2.5,3.3,4,5,5.5]
#IPW
IPW = np.linspace(0,15,16,dtype=int)
int_array = np.arange(16)
hex_array = np.array([f"0x{hex(i)[2:].upper()}" for i in int_array])
# hex_array = np.array([hex(i) for i in int_array]) # 使用列表推导式创建


##### XTHF 频率和CFG的关系 #####
#从excel文件中提取具体目录
# wb = openpyxl.load_workbook('475XTHF.xlsx')
# sheet = sheet = wb['FRE']
# data_array = []

# for M in M_Fre:
#     # 读取5行5列的数据
#     for row in range(2 + M*17,18+M*17):  # 从第1行到第17行
#         row_data = []
#         for col in range(2, 7):  # 从第1列到第5列2，7
#             cell = sheet.cell(row=row, column=col)
#             row_data.append(cell.value)
#         data_array.append(row_data)

#     # 将列表转换为NumPy数组
#     data_array = np.array(data_array)
#     print(data_array)

#     #将数组中频率计算后转化为PPM
#     # transformed_array = transform_array(data_array,32768, 32768)

#     data_array = data_array.T
#     transformed_array = substract_column(data_array,8)

#     # transformed_array = transformed_array.T
#     print(transformed_array)

#     #开始画图
#     plt.figure(figsize=(12, 8))
#     for i in range(transformed_array.shape[0]):
#         plt.plot(IPW,transformed_array[i], label=f'FM475#{i+1}')
#         plt.scatter(IPW,transformed_array[i])

#     plt.grid(True)
#     # 设置图表标题和坐标轴标签
#     if M  != 5:
#         plt.title(f'{CryFre[M]}M晶振频率随CFG的变化 @VDD={3.3}V',fontsize = 24)
#     else:
#         plt.title(f'{CryFre[M]}M陶振频率随CFG的变化 @VDD={3.3}V',fontsize = 24)
#     plt.xticks(IPW,hex_array,fontsize = 18) ##plt.xticks(IPW,hex_array,fontsize = 18) 
#     plt.yticks(fontsize=18)
#     plt.xlabel('CFG',fontsize = 24)
#     plt.ylabel('PPM',fontsize = 24)
#     plt.legend(fontsize = 'large')
#     # 设置图形的dpi
#     plt.savefig(f'#XTHF_FRE_{M}M.png', dpi=300,bbox_inches = 'tight')  # 保存图片时设置dpi
#     # # 显示图表
# #    plt.show()

#     cell.value = None
#     data_array = []
##### XTHF 频率和CFG的关系 #####


##### XTHF 频率和VDD的关系 #####
# wb = openpyxl.load_workbook('475XTHF.xlsx')
# sheet = sheet = wb['FRE_VPP']
# data_array = []

# for M in M_Fre:
#     # 读取5行5列的数据
#     for row in range(2 + M*7,7+M*7):  # 从第1行到第17行
#         row_data = []
#         for col in range(2, 7):  # 从第1列到第5列2，7
#             cell = sheet.cell(row=row, column=col)
#             row_data.append(cell.value)
#         data_array.append(row_data)

#     # 将列表转换为NumPy数组
#     data_array = np.array(data_array)
#     print(data_array)

#     #将数组中频率计算后转化为PPM
#     # transformed_array = transform_array(data_array,32768, 32768)

#     data_array = data_array.T
#     transformed_array = substract_column(data_array,2)

#     # transformed_array = transformed_array.T
#     print(transformed_array)

#     #开始画图
#     plt.figure(figsize=(12, 8))
#     for i in range(transformed_array.shape[0]):
#         plt.plot(VDD,transformed_array[i], label=f'FM475#{i+1}')
#         plt.scatter(VDD,transformed_array[i])

#     plt.grid(True)
#     # 设置图表标题和坐标轴标签
#     if M  != 5:
#         plt.title(f'{CryFre[M]}M晶振频率随VDD的变化 @CFG=0x8',fontsize = 24)
#     else:
#         plt.title(f'{CryFre[M]}M陶振频率随VDD的变化 @CFG=0x8',fontsize = 24)
#     plt.xticks(fontsize = 18) ##plt.xticks(IPW,hex_array,fontsize = 18) 
#     plt.yticks(fontsize=18)
#     plt.xlabel('VDD (V)',fontsize = 24)
#     plt.ylabel('PPM',fontsize = 24)
#     plt.legend(fontsize = 'large')
#     # 设置图形的dpi
#     plt.savefig(f'#XTHF_VDD_{M}M.png', dpi=300,bbox_inches = 'tight')  # 保存图片时设置dpi
#     # # 显示图表
#    # plt.show()

#     cell.value = None
#     data_array = []
##### XTHF 频率和VDD的关系 #####


##### XTHF功率 ######
# #从excel文件中提取具体目录
# wb = openpyxl.load_workbook('475XTHF.xlsx')
# sheet = sheet = wb['Power']
# data_array = []

# ## 3.3V ##
# for M in M_Fre:
#     # 读取15行5列的数据
#     for row in range(2 + M*18,18+M*18):  # 从第1行到第17行
#         row_data = []
#         for col in range(2, 7):  # 从第1列到第5列2，7
#             cell = sheet.cell(row=row, column=col)
#             row_data.append(cell.value)
#         data_array.append(row_data)

#     # 将列表转换为NumPy数组
#     data_array = np.array(data_array)
#     print(data_array)

#     transformed_array = data_array.T
#     # transformed_array = transformed_array.T
#     print(transformed_array)

#     #开始画图
#     plt.figure(figsize=(12, 8))
#     for i in range(transformed_array.shape[0]):
#         plt.plot(IPW,transformed_array[i], label=f'FM475#{i+1}')
#         plt.scatter(IPW,transformed_array[i])

#     plt.grid(True)
#     # 设置图表标题和坐标轴标签
#     if M  != 5:
#         plt.title(f'{CryFre[M]}M晶振功耗随CFG的变化 @VDD={3.3}V',fontsize = 24)
#     else:
#         plt.title(f'{CryFre[M]}M陶振功耗随CFG的变化 @VDD={3.3}V',fontsize = 24)
#     plt.xticks(IPW,hex_array,fontsize = 18) ##plt.xticks(IPW,hex_array,fontsize = 18) 
#     plt.yticks(fontsize=18)
#     plt.xlabel('CFG',fontsize = 24)
#     plt.ylabel('功耗 (uA)',fontsize = 24)
#     plt.legend(fontsize = 'large')
#     # 设置图形的dpi
#     plt.savefig(f'#XTHF_Power_{M}M.png', dpi=300,bbox_inches = 'tight')  # 保存图片时设置dpi
#     # # 显示图表
#     #plt.show()

#     cell.value = None
#     data_array = []

# ## 5V##
# for M in M_Fre:
#     # 读取15行5列的数据
#     for row in range(2 + M*18,18+M*18):  # 从第1行到第17行
#         row_data = []
#         for col in range(9, 14):  # 列数
#             cell = sheet.cell(row=row, column=col)
#             row_data.append(cell.value)
#         data_array.append(row_data)

#     # 将列表转换为NumPy数组
#     data_array = np.array(data_array)
#     print(data_array)

#     transformed_array = data_array.T
#     # transformed_array = transformed_array.T
#     print(transformed_array)

#     #开始画图
#     plt.figure(figsize=(12, 8))
#     for i in range(transformed_array.shape[0]):
#         plt.plot(IPW,transformed_array[i], label=f'FM475#{i+1}')
#         plt.scatter(IPW,transformed_array[i])

#     plt.grid(True)
#     # 设置图表标题和坐标轴标签
#     if M  != 5:
#         plt.title(f'{CryFre[M]}M晶振功耗随CFG的变化 @VDD={5}V',fontsize = 24)
#     else:
#         plt.title(f'{CryFre[M]}M陶振功耗随CFG的变化 @VDD={5}V',fontsize = 24)
#     plt.xticks(IPW,hex_array,fontsize = 18) ##plt.xticks(IPW,hex_array,fontsize = 18) 
#     plt.yticks(fontsize=18)
#     plt.xlabel('CFG',fontsize = 24)
#     plt.ylabel('功耗 (uA)',fontsize = 24)
#     plt.legend(fontsize = 'large')
#     # 设置图形的dpi
#     plt.savefig(f'#XTHF_Power2_{M}M.png', dpi=300)  # 保存图片时设置dpi
#     # # 显示图表
#     #plt.show()

#     cell.value = None
#     data_array = []
##### XTHF功率 ######


# #### XTHF振幅 ####
# #从excel文件中提取具体目录
# wb = openpyxl.load_workbook('475XTHF.xlsx')
# sheet = sheet = wb['VPP']
# data_array = []

# ## 3.3V OUT端 ##
# for M in M_Fre:
#     # 读取5行5列的数据
#     for row in range(2 + 18*M,18 + 18*M):  # 从第1行到第17行
#         row_data = []
#         for col in range(2, 7):  # 从第1列到第5列2，7
#             cell = sheet.cell(row=row, column=col)
#             row_data.append(cell.value)
#         data_array.append(row_data)

#     # 将列表转换为NumPy数组
#     data_array = np.array(data_array)
#     print(data_array)


#     # transformed_array = data_array - data_array[15]
#     transformed_array = data_array.T
#     # transformed_array = transformed_array.T
#     print(transformed_array)

#     #开始画图
#     plt.figure(figsize=(12, 8))
#     for i in range(transformed_array.shape[0]):
#         plt.plot(IPW,transformed_array[i], label=f'FM475#{i+1}')
#         plt.scatter(IPW,transformed_array[i])

#     plt.grid(True)
#     # 设置图表标题和坐标轴标签
#     if M  != 5:
#         plt.title(f'{CryFre[M]}M晶振OUT端振幅随CFG的变化 @VDD={3.3}V',fontsize = 24)
#     else:
#         plt.title(f'{CryFre[M]}M陶振OUT端振幅随CFG的变化 @VDD={3.3}V',fontsize = 24)
#     plt.xticks(IPW,hex_array,fontsize = 18) ##plt.xticks(IPW,hex_array,fontsize = 18) 
#     plt.yticks(fontsize=18)
#     plt.xlabel('CFG',fontsize = 24)
#     plt.ylabel('VPP (mV)',fontsize = 24)
#     plt.legend(fontsize = 'large')
#     # 设置图形的dpi
#     plt.savefig(f'#XTHF_OUT_{M}M.png', dpi=300)  # 保存图片时设置dpi
#     # # 显示图表
#     # plt.show()

#     cell.value = None
#     data_array = []

# ## 3.3V IN端 ##
# for M in M_Fre:
#     # 读取5行5列的数据
#     for row in range(2 + 18*M,18 + 18*M):  # 从第1行到第17行
#         row_data = []
#         for col in range(8, 13):  # 从第9列到第13列
#             cell = sheet.cell(row=row, column=col)
#             row_data.append(cell.value)
#         data_array.append(row_data)

#     # 将列表转换为NumPy数组
#     data_array = np.array(data_array)
#     print(data_array)


#     # transformed_array = data_array - data_array[15]
#     transformed_array = data_array.T
#     # transformed_array = transformed_array.T
#     print(transformed_array)

#     #开始画图
#     plt.figure(figsize=(12, 8))
#     for i in range(transformed_array.shape[0]):
#         plt.plot(IPW,transformed_array[i], label=f'FM475#{i+1}')
#         plt.scatter(IPW,transformed_array[i])

#     plt.grid(True)
#     # 设置图表标题和坐标轴标签
#     if M  != 5:
#         plt.title(f'{CryFre[M]}M晶振IN端振幅随CFG的变化 @VDD={3.3}V',fontsize = 24)
#     else:
#         plt.title(f'{CryFre[M]}M陶振IN端振幅随CFG的变化 @VDD={3.3}V',fontsize = 24)
#     plt.xticks(IPW,hex_array,fontsize = 18) ##plt.xticks(IPW,hex_array,fontsize = 18) 
#     plt.yticks(fontsize=18)
#     plt.xlabel('CFG',fontsize = 24)
#     plt.ylabel('VPP (mV)',fontsize = 24)
#     plt.legend(fontsize = 'large')
#     # 设置图形的dpi
#     plt.savefig(f'#XTHF_IN_{M}M.png', dpi=300)  # 保存图片时设置dpi
#     # # 显示图表
#     # plt.show()

#     cell.value = None
#     data_array = []

# ## 5V OUT端 ##
# for M in M_Fre:
#     # 读取5行5列的数据
#     for row in range(2 + 18*M,18 + 18*M):  # 从第1行到第17行
#         row_data = []
#         for col in range(14, 19):  # 从第1列到第5列2，7
#             cell = sheet.cell(row=row, column=col)
#             row_data.append(cell.value)
#         data_array.append(row_data)

#     # 将列表转换为NumPy数组
#     data_array = np.array(data_array)
#     print(data_array)


#     # transformed_array = data_array - data_array[15]
#     transformed_array = data_array.T
#     # transformed_array = transformed_array.T
#     print(transformed_array)

#     #开始画图
#     plt.figure(figsize=(12, 8))
#     for i in range(transformed_array.shape[0]):
#         plt.plot(IPW,transformed_array[i], label=f'FM475#{i+1}')
#         plt.scatter(IPW,transformed_array[i])

#     plt.grid(True)
#     # 设置图表标题和坐标轴标签
#     if M  != 5:
#         plt.title(f'{CryFre[M]}M晶振OUT端振幅随CFG的变化 @VDD={5}V',fontsize = 24)
#     else:
#         plt.title(f'{CryFre[M]}M陶振OUT端振幅随CFG的变化 @VDD={5}V',fontsize = 24)
#     plt.xticks(IPW,hex_array,fontsize = 18) ##plt.xticks(IPW,hex_array,fontsize = 18) 
#     plt.yticks(fontsize=18)
#     plt.xlabel('CFG',fontsize = 24)
#     plt.ylabel('VPP (mV)',fontsize = 24)
#     plt.legend(fontsize = 'large')
#     # 设置图形的dpi
#     plt.savefig(f'#XTHF_OUT2_{M}M.png', dpi=300)  # 保存图片时设置dpi
#     # # 显示图表
#     # plt.show()

#     cell.value = None
#     data_array = []

# ## 3.3V IN端 ##
# for M in M_Fre:
#     # 读取5行5列的数据
#     for row in range(2 + 18*M,18 + 18*M):  # 从第1行到第17行
#         row_data = []
#         for col in range(20, 25):  # 从第9列到第13列
#             cell = sheet.cell(row=row, column=col)
#             row_data.append(cell.value)
#         data_array.append(row_data)

#     # 将列表转换为NumPy数组
#     data_array = np.array(data_array)
#     print(data_array)


#     # transformed_array = data_array - data_array[15]
#     transformed_array = data_array.T
#     # transformed_array = transformed_array.T
#     print(transformed_array)

#     #开始画图
#     plt.figure(figsize=(12, 8))
#     for i in range(transformed_array.shape[0]):
#         plt.plot(IPW,transformed_array[i], label=f'FM475#{i+1}')
#         plt.scatter(IPW,transformed_array[i])

#     plt.grid(True)
#     # 设置图表标题和坐标轴标签
#     if M  != 5:
#         plt.title(f'{CryFre[M]}M晶振IN端振幅随CFG的变化 @VDD={5}V',fontsize = 24)
#     else:
#         plt.title(f'{CryFre[M]}M陶振IN端振幅随CFG的变化 @VDD={5}V',fontsize = 24)
#     plt.xticks(IPW,hex_array,fontsize = 18) ##plt.xticks(IPW,hex_array,fontsize = 18) 
#     plt.yticks(fontsize=18)
#     plt.xlabel('CFG',fontsize = 24)
#     plt.ylabel('VPP (mV)',fontsize = 24)
#     plt.legend(fontsize = 'large')
#     # 设置图形的dpi
#     plt.savefig(f'#XTHF_IN1_{M}M.png', dpi=300)  # 保存图片时设置dpi
#     # # 显示图表
#     # plt.show()

#     cell.value = None
#     data_array = []
# #### XTHF振幅 ####

#### XTHF建立时间 ####
# #从excel文件中提取具体目录
# wb = openpyxl.load_workbook('475XTHF.xlsx')
# sheet = sheet = wb['TIEM']
# data_array = []

# ## 3.3V ##
# for M in M_Fre:
#     # 读取16行5列的数据
#     for row in range(2 + M*17,18+M*17):  # 从第1行到第17行
#         row_data = []
#         for col in range(2, 7):  # 从第1列到第5列2，7
#             cell = sheet.cell(row=row, column=col)
#             row_data.append(cell.value)
#         data_array.append(row_data)

#     # 将列表转换为NumPy数组
#     data_array = np.array(data_array)
#     print(data_array)

#     transformed_array = data_array.T
#     # transformed_array = transformed_array.T
#     print(transformed_array)

#     #开始画图
#     plt.figure(figsize=(12, 8))

#     # 绘制主图
#     ax = plt.gca()  # 获取当前轴
#     for i in range(transformed_array.shape[0]):
#         ax.plot(IPW,transformed_array[i], label=f'FM475#{i+1}')
#         ax.scatter(IPW,transformed_array[i])

#     ax.legend()

#     if M != 5:
#         # 创建局部放大子图（放在右下角）
#         axins = ax.inset_axes([0.55, 0.45, 0.4, 0.4])  # 调整位置和大
#         # axins = ax.inset_axes([0.6, 0.5, 0.35, 0.35])  # [左, 下, 宽, 高] 相对比例

#         # 计算右半边数据的索引
#         right_half_index = 8

#         # 绘制局部放大图
#         for i in range(transformed_array.shape[0]):
#             axins.plot(IPW[right_half_index:], transformed_array[i][right_half_index:])
#             axins. scatter(IPW[right_half_index:], transformed_array[i][right_half_index:])

#         # # 设置局部放大区域的显示范围（可以根据你的数据调整）
#         # x1, x2 = IPW[right_half_index], IPW[-1]
#         # y1, y2 = np.min(transformed_array[:, right_half_index:]), np.max(transformed_array[:, right_half_index:])

#     # 设置局部放大区域的显示范围（增加空白）
#         x1, x2 = IPW[right_half_index] - (IPW[-1] - IPW[right_half_index]) * 0.1, IPW[-1] + (IPW[-1] - IPW[right_half_index]) * 0.1 #X轴增加空白
#         y1, y2 = np.min(transformed_array[:, right_half_index:]) - (np.max(transformed_array[:, right_half_index:])-np.min(transformed_array[:, right_half_index:])) * 0.1, np.max(transformed_array[:, right_half_index:]) + (np.max(transformed_array[:, right_half_index:])-np.min(transformed_array[:, right_half_index:])) * 0.1 #Y轴增加空白
    

#         axins.set_xlim(x1, x2)
#         axins.set_ylim(y1, y2)

#         # 在主图中绘制局部放大区域的边框
#         ax.indicate_inset_zoom(axins, edgecolor="red")
    
#     plt.grid(True)
#     # 设置图表标题和坐标轴标签
#     if M  != 5:
#         plt.title(f'{CryFre[M]}M晶振建立时间随CFG的变化 @VDD={3.3}V',fontsize = 24)
#     else:
#         plt.title(f'{CryFre[M]}M陶振建立时间随CFG的变化 @VDD={3.3}V',fontsize = 24)
#     plt.xticks(IPW,hex_array,fontsize = 18) ##plt.xticks(IPW,hex_array,fontsize = 18) 
#     plt.yticks(fontsize=18)
#     # plt.yscale('log')
#     plt.xlabel('CFG',fontsize = 24)
#     plt.ylabel('建立时间 (us)',fontsize = 24)
#     plt.legend(fontsize = 'large')
#     # 设置图形的dpi
#     plt.savefig(f'#XTHF_Time_{M}M.png', dpi=300)  # 保存图片时设置dpi
#     # # 显示图表
#     #plt.show()

#     cell.value = None
#     data_array = []


#### 将数据转化为word中表格 ####

#从excel文件中提取具体目录
wb = openpyxl.load_workbook('475XTHF.xlsx')
sheet = sheet = wb['TIEM']
data_array = []

for M in M_Fre:
    append_excel_to_word_table(
        '475XTHF.xlsx',  # Excel 文件路径
        '475XTHF.docx',  # Word 文件保存路径
        sheet_name='FRE1',  # 工作表名称
        start_row=0+M*17,  # 起始行（第 3 行）
        end_row=16+M*17,  # 结束行（第 7 行，不包含）
        start_col=0,  # 起始列（第 2 列）
        end_col=6   # 结束列（第 4 列，不包含）
    )

for M in M_Fre:
    append_excel_to_word_table(
        '475XTHF.xlsx',  # Excel 文件路径
        '475XTHF.docx',  # Word 文件保存路径
        sheet_name='TIEM',  # 工作表名称
        start_row=0+M*17,  # 起始行（第 3 行）
        end_row=16+M*17,  # 结束行（第 7 行，不包含）
        start_col=0,  # 起始列（第 2 列）
        end_col=6   # 结束列（第 4 列，不包含）
    )

#### XTHF建立时间 ####