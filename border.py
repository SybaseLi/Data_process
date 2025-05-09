import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH# 导入 WD_ALIGN_PARAGRAPH 枚举，用于设置段落对齐方式

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Pt
from docx.oxml.ns import qn

def set_boeder(elm,border_name,sz,color,val):
    border = elm.find(qn(f'w:{border_name}'))
    if border is None:
        border = elm.makeelement(qn(f'w:{border_name}'))
        elm.append(border)
    
    border.set(qn('w:sz'),str(sz))
    border.set(qn('w:color'),color)
    border.set(qn('w:val'),val)


# 检查 Word 文件是否存在，如果存在则读取，否则创建新文档
if os.path.exists('XTHF.docx'):
    doc = Document('XTHF.docx')
else:
    doc = Document()

table = doc.add_table(6, 8)

num_rows = len(table.rows)
num_cols = len(table.columns)
for i in range(num_rows):
    for j in range(num_cols):
        cell = table.cell(i,j)
        tc = cell._element
        # 检查tcBorders是否存在
        tcBorders_list = tc.xpath('.//w:tcBorders')
        if not tcBorders_list:
            # 如果不存在，创建tcBorders元素
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        else:
            tcBorders = tcBorders_list[0]

        #处理顶部边框
        if i == 0:
            set_boeder(tcBorders,'top',4,'000000','double')
            top_inner = tcBorders.makeelement(qn('w:top'),{
                qn('w:sz'):'4',
                qn('w:color'):'000000',
                qn('w:val'):'double'
            }
            )
            tcBorders.append(top_inner)
        else:
            set_boeder(tcBorders,'top',4,'000000','single')
            # #非顶部单元格
            # top_border = tcBorders.find(qn('w:top'))
            # if top_border is not None:
            #     tcBorders.remove(top_border)
    
        #处理左侧边框
        if j == 0:
            set_boeder(tcBorders,'left',4,'000000','double')
            left_inner = tcBorders.makeelement(qn('w:left'),{
                qn('w:sz'):'4',
                qn('w:color'):'000000',
                qn('w:val'):'double'
            }
            )
            tcBorders.append(left_inner)
        else:
            set_boeder(tcBorders,'left',4,'000000','single')
        
        #处理右侧边框
        if j == num_cols-1:
            set_boeder(tcBorders,'right',4,'000000','double')
            right_inner = tcBorders.makeelement(qn('w:right'),{
                qn('w:sz'):'4',
                qn('w:color'):'000000',
                qn('w:val'):'double'
            }
            )
            tcBorders.append(right_inner)
        else:
            set_boeder(tcBorders,'right',4,'000000','single')

        #处理底部边框
        if i == num_rows-1:
            set_boeder(tcBorders,'bottom',4,'000000','double')
            bottom_inner = tcBorders.makeelement(qn('w:bottom'),{
                qn('w:sz'):'4',
                qn('w:color'):'000000',
                qn('w:val'):'double'                }
            )
            tcBorders.append(bottom_inner)
        else:
            set_boeder(tcBorders,'bottom',4,'000000','single')

    try:
        doc.save('XTHF.docx')
        print(f"成功将 Excel 数据追加到 Word 表格：{'XTHF.docx'}")
    except Exception as e:
        print(f"保存 Word 文件时出错：{e}")